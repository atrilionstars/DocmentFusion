import os
import argparse
import json
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
from typing import List, Tuple
from docx.shared import Pt


BASE_URL = "http://192.168.1.13:8000/v1"
API_KEY = "2JysWWdHfyKvp2AsGYznw7pwPfkwDehtPZHEtj26GIA"
MODEL_NAME = "qwen-max"

# 请求头
headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {API_KEY}"
}

def read_docx(file_path: str) -> str:
    """读取docx文档内容"""
    # 创建Document对象读取Word文档
    doc = Document(file_path)
    full_text = []
    # 遍历文档中的每个段落，提取文本内容
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 将所有段落文本用换行符连接成一个字符串并返回
    return '\n'.join(full_text)
def create_prompt(files_content: List[Tuple[str, str]], instruction: str = "") -> str:
    """创建提示词"""
    # 初始化提示词，介绍后续是多个文档的内容
    prompt = "以下是多个文档的内容:\n\n"

    # 遍历每个文档，将文档名和内容添加到提示词中
    for i, (filename, content) in enumerate(files_content):
        prompt += f"文档 {i+1} ({filename}):\n{content}\n\n"

    # 如果提供了额外说明，将其添加到提示词中
    if instruction:
        prompt += f"额外说明: {instruction}\n\n"

    # 添加最后的指令，要求模型对文档内容进行融合重写
    prompt += "根据提交的内容，写字100字的内容。"
    return prompt

def call_llm(prompt: str) -> str:
    """调用LLM API生成内容"""

    url = f"{BASE_URL}/chat/completions"
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是一个专业的文档融合专家。请根据用户提供的多个文档内容，进行融合重写，确保内容连贯、逻辑清晰，并且保留原文的核心信息。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 4000
    }

    try:
        print(f"发送请求到API，提示词长度: {len(prompt)} 字符")
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        result = response.json()["choices"][0]["message"]["content"].strip()

        # 调试输出
        print(f"API返回内容长度: {len(result)} 字符")
        if len(result) < 10:  # 如果内容非常短，可能有问题
            print(f"警告: API返回的内容异常短: {result[:50]}...")

        return result
    except Exception as e:
        print(f"API调用错误: {e}")
        return None

def save_to_docx(content: str, output_file: str) -> None:
    """保存内容到docx文档，解决内容不可见及NameError问题"""
    if not content.strip():
        print("错误: 要保存的内容为空")
        return

    print(f"保存前内容检查 - 前100个字符: {content[:100]}")
    print(f"保存前内容检查 - 后100个字符: {content[-100:]}")

    doc = Document()
    title = doc.add_heading("融合重写文档", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    from datetime import datetime
    date_str = datetime.now().strftime("%Y年%m月%d日")
    date_para = doc.add_paragraph(date_str)
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 移除可能导致显示问题的不可见字符
    from unicodedata import category
    cleaned_content = ''.join(c for c in content if category(c)[0] != 'C')

    # 按段落分割，处理不同的换行符情况
    paragraphs = cleaned_content.split('\n\n')
    paragraphs = [p for p in paragraphs if p.strip()]

    if not paragraphs:
        print("警告: 处理后没有有效段落")
        paragraphs = [cleaned_content]  # 作为最后手段，使用完整内容

    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            # 创建段落并添加运行，设置基本字体属性
            p = doc.add_paragraph()
            run = p.add_run(para_text)

            # 设置基本字体属性，避免可能的隐藏格式
            run.font.name = '宋体'  # 设置中文字体
            run.font.size = Pt(12)  # 修正：直接使用Pt（已导入）
            run.font.hidden = False  # 确保文本不被隐藏

    doc.save(output_file)
    print(f"已保存融合文档到: {output_file}，内容长度: {len(content)} 字符")

    # 保存一个文本版本用于对比
    text_output_file = output_file.replace('.docx', '_debug.txt')
    with open(text_output_file, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"已保存调试文本版本到: {text_output_file}")
def main():
    # 创建命令行参数解析器，设置工具描述
    parser = argparse.ArgumentParser(description="多文档融合重写工具")
    # 添加输入文件参数 - 允许多个docx文件路径
    parser.add_argument("-i", "--input", nargs="+", help="输入的docx文件路径")
    # 添加输出文件参数 - 指定融合后的文档保存路径
    parser.add_argument("-o", "--output", help="输出的docx文件路径")
    # 添加额外说明参数 - 可以提供给模型的额外指令
    parser.add_argument("-t", "--instruction", help="额外说明或指令")
    # 解析命令行参数
    args = parser.parse_args()
    # 处理默认输入文件 - 如果未指定输入文件，使用当前目录下的所有docx文件
    if not args.input:
        print("未指定输入文件，将使用当前目录下的所有docx文件")
        current_dir = os.getcwd()
        args.input = [os.path.join(current_dir, f) for f in os.listdir(current_dir) if f.lower().endswith('.docx')]

        # 检查是否找到docx文件
        if not args.input:
            print("错误: 当前目录下没有找到docx文件")
            return

    # 处理默认输出文件 - 如果未指定输出文件，使用默认名称
    if not args.output:
        print("未指定输出文件，将使用默认名称")
        current_dir = os.getcwd()
        args.output = os.path.join(current_dir, "融合文档_自动生成.docx")

    # 检查输入文件是否存在 - 收集不存在的文件
    invalid_files = [f for f in args.input if not os.path.exists(f)]
    if invalid_files:
        print(f"错误: 以下文件不存在: {', '.join(invalid_files)}")
        return

    # 检查输出文件目录是否存在 - 如果不存在则创建
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 读取所有文档内容 - 使用进度条显示读取进度
    print("正在读取文档...")
    files_content = []
    for file_path in tqdm(args.input):
        content = read_docx(file_path)
        files_content.append((os.path.basename(file_path), content))

    # 创建提示词 - 准备提供给LLM的输入
    print("正在准备提示词...")
    prompt = create_prompt(files_content, args.instruction)

    # 调用LLM进行融合重写 - 调用API并显示处理状态
    print("正在调用LLM进行融合重写...")
    result = call_llm(prompt)

    # 处理API返回结果 - 如果成功则保存文档，否则显示错误信息
    if result:
        # 保存结果到docx
        save_to_docx(result, args.output)
    else:
        print("融合失败，未能获取有效内容")

if __name__ == "__main__":
    # 程序入口点 - 调用主函数
    main()